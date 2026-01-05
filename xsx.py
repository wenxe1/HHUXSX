import os
KEY = 1 
import io
import re
import random
from docx import Document  # 读写word
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
import streamlit as st

st.set_page_config(
    page_title="HHU习思想", 
    page_icon="📖",  # 使用一个书本图标
    layout="centered"  # 使用居中布局，更适合阅读
)
# 注入自定义CSS
st.markdown("""
<style>
[data-testid="stDecoration"] { display: none; }
[data-testid="stHeader"] { background: none; box-shadow: none; }
[data-testid="stSidebar"] div[data-baseweb="notification"] { width: fit-content; display: inline-block; margin: 2px 0 6px 0; }
html, body, [data-testid="stAppViewContainer"], .stApp { caret-color: transparent; }
[data-testid="stSidebar"] [data-testid="stColumns"] > div { padding-left: 0 !important; padding-right: 0 !important; }
[data-testid="stSidebar"] [data-testid="stColumns"] { gap: 0 !important; }
[data-testid="stSidebar"] [data-testid="stDownloadButton"] { display: inline-block; margin: 0; }
.stSidebar { font-size: 13px; }
[data-testid="stSidebar"] hr { margin-top: 6px; margin-bottom: 6px; }
[data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 { margin-top: 6px; margin-bottom: 4px; }
.result-ok { background: #e8f5e9; color: #1b5e20; padding: 10px 12px; border-radius: 8px; margin: 6px 0 14px 0; }
.result-bad { background: #fdecea; color: #b71c1c; padding: 10px 12px; border-radius: 8px; margin: 6px 0 14px 0; }
html, body { font-size: 14px; }
</style>
""", unsafe_allow_html=True)


def load_questions(filepath):
    """加载题库"""
    # 解密题库wrod文档
    with open(filepath, 'rb') as f:
        file_data = bytearray(f.read())
    for i in range(len(file_data)):
        file_data[i] ^= KEY
    file_stream = io.BytesIO(file_data)
    doc = Document(file_stream)

    questions = []
    current_q = None
    
    # 正则表达式匹配规则
    q_pattern = re.compile(r'^\s*(\d+)[\.、．]\s*(.*)') 
    opt_pattern = re.compile(r'^\s*([A-EＡ-Ｅ])[\s\.、．\)）]\s*(.*)')
    ans_pattern = re.compile(r'答案\s*[:：]\s*([A-EＡ-Ｅ]+)')
    
    for para in doc.paragraphs:
        lines = para.text.splitlines()

        # 遍历分割出的每一行
        for line in lines:
            text = line.strip()

            # 跳过空行
            if not text:
                continue  
        
            # 检查是否是题目
            q_match = q_pattern.match(text)
            if q_match:
                if current_q:  # 保存上一题
                    if current_q.get("options"): 
                        questions.append(current_q)
                current_q = {
                    "id": q_match.group(1).strip(),
                    "title": q_match.group(2).strip(),
                    "options": [],
                    "answer": "",
                    "type": "单选" # 默认单选 后面根据答案长度修正
                }
                continue
            
            # 检查是否是选项
            opt_match = opt_pattern.match(text)
            if current_q and opt_match:
                option_letter = opt_match.group(1).strip()
                option_content = opt_match.group(2).strip()
                full_option = f"{option_letter}. {option_content}"
                current_q["options"].append(full_option)
                continue
                
            # 检查是否是答案
            ans_match = ans_pattern.search(text)
            if current_q and ans_match:
                ans = ans_match.group(1).strip()
                ans = ans.replace('Ａ', 'A').replace('Ｂ', 'B').replace('Ｃ', 'C').replace('Ｄ', 'D').replace('Ｅ', 'E')
                current_q["answer"] = ans
                if len(ans) > 1:
                    current_q["type"] = "多选"
    
    if current_q:  # 加入最后一题
        questions.append(current_q)
                
    return questions

def output_wrong_doc(wrong_questions, with_answer=False):
    """生成错题文档"""
    doc = Document()
    
    # 修改字体样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for idx, q in enumerate(wrong_questions):
        # 写入题目
        p = doc.add_paragraph()
        run = p.add_run(f"{q['id']}. [{q['type']}] {q['title']}")
        # run.bold = True
        
        # 写入选项
        for opt in q['options']:
            # 提取选项字母，例如从 "A. 北京" 中提取 "A"
            option_letter = opt.split('.')[0].strip()
            
            # 检查当前选项的字母是否在正确答案中
            if with_answer and option_letter in q['answer']:
                # 正确答案选项标黄
                p_opt = doc.add_paragraph()
                run_opt = p_opt.add_run(opt)
                run_opt.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                # 不是正确答案选项
                doc.add_paragraph(opt)
            
        doc.add_paragraph("")  # 空行分隔
        
    # 保存到内存流
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    return f


# 初始化SessionState 类似于全局变量 用于存储刷题状态
if 'mistakes' not in st.session_state:
    st.session_state['mistakes'] = [] 
if 'current_q' not in st.session_state:
    st.session_state['current_q'] = None
if 'quiz_data' not in st.session_state:
    st.session_state['quiz_data'] = []
if 'user_choice' not in st.session_state:
    st.session_state['user_choice'] = None
if 'submitted' not in st.session_state:
    st.session_state['submitted'] = False
if 'last_correct' not in st.session_state:
    st.session_state['last_correct'] = None

# --- 侧边栏 ---
with st.sidebar:
    st.title("⚙️ 功能菜单")

    base_path = os.path.dirname(os.path.abspath(__file__))
    QUESTIONS_DIR = os.path.join(base_path, "questions")

    # 2. 查找所有章节文件并让用户选择
    try:
        # os.listdir() 获取文件夹下所有文件名
        # [f for ... if f.endswith('.docx')] 筛选出word文档
        # sorted() 对文件名进行排序
        def chapter_sort_key(name):
            m = re.search(r'(\d+)', name)
            return int(m.group(1)) if m else float('inf')
        available_chapters = sorted(
            [f for f in os.listdir(QUESTIONS_DIR) if f.endswith('.docx')],
            key=chapter_sort_key
        )
        
        def chapter_label(name):
            m = re.search(r'ch(\d+)\.docx', name, re.IGNORECASE)
            if m:
                n = int(m.group(1))
                return "导论" if n == 0 else f"第{n}章"
            return name
        
        selected_chapters = st.multiselect(
            "选择要练习的章节:",
            options=available_chapters,
            default=available_chapters,
            format_func=chapter_label
        )
        
        # 3. 添加一个按钮来触发加载
        if st.button("🚀 加载选中章节"):
            if not selected_chapters:
                st.warning("请至少选择一个章节！")
            else:
                # 清空旧数据
                st.session_state['quiz_data'] = []
                all_questions = []
                
                # 4. 循环读取和合并数据
                for chapter_file in selected_chapters:
                    # os.path.join() 用来拼接完整的文件路径
                    filepath = os.path.join(QUESTIONS_DIR, chapter_file)
                    try:
                        questions_from_file = load_questions(filepath)
                        all_questions.extend(questions_from_file) # 使用 extend 合并列表
                    except Exception as e:
                        st.error(f"解析文件 {chapter_file} 失败: {e}")
                
                st.session_state['quiz_data'] = all_questions
                # 重置刷题状态
                st.session_state['current_q'] = None
                st.session_state['submitted'] = False
                
                st.success(f"成功加载 {len(all_questions)} 道题！")
                st.rerun() # 立即刷新界面以开始答题

    except FileNotFoundError:
        st.error(f"错误：未找到题库文件夹 '{QUESTIONS_DIR}'。")
        st.info(f"请确保在程序同级目录下有一个名为 '{QUESTIONS_DIR}' 的文件夹，并将题库文件放入其中。")


    st.divider()
    st.subheader("📥 错题管理")
    if len(st.session_state['mistakes']) > 0:
        doc_pure = output_wrong_doc(st.session_state['mistakes'], with_answer=False)
        st.download_button("导出刷题版", doc_pure, "错题刷题版.docx")
        doc_ans = output_wrong_doc(st.session_state['mistakes'], with_answer=True)
        st.download_button("导出复习版", doc_ans, "错题复习版.docx")
    else:
        st.info("暂无错题可导出")
    if st.button("🗑️ 清空错题记录"):
        st.session_state['mistakes'] = []
        st.rerun()

# --- 主界面 ---
st.title("HHU习思想")

if not st.session_state.get('quiz_data'):
    st.info("👈 请在左侧选择章节并点击“加载选中章节”开始刷题")
else:
    # 抽题逻辑 (使用之前修复过的安全版本)
    if st.session_state.get('current_q') is None:
        if not st.session_state['quiz_data']:
            st.warning("🎉 恭喜！您已完成所有题目。")
        else:
            valid_question_found = False
            while not valid_question_found and st.session_state['quiz_data']:
                q_candidate = random.choice(st.session_state['quiz_data'])
                if q_candidate and q_candidate.get("options"):
                    st.session_state['current_q'] = q_candidate
                    st.session_state['submitted'] = False
                    st.session_state['user_choice'] = None
                    valid_question_found = True
                else:
                    st.warning(f"警告: 题目ID {q_candidate.get('id', 'N/A')} 因缺少选项被跳过。")
                    st.session_state['quiz_data'].remove(q_candidate)
            if not valid_question_found:
                 st.error("所有题目都因格式问题被跳过，无法出题。")
                 st.stop()

    if st.session_state.get('current_q'):
        q = st.session_state['current_q']
        
        st.markdown(f"#### {q['id']}. [{q['type']}] {q['title']}")
        st.divider()
        
        choice = None
        
        # --- 核心修改：统一使用 Checkbox ---
        
        # 1. 初始化一个字典来存储每个选项的选中状态
        if 'option_states' not in st.session_state or st.session_state.get('current_q_id') != q['id']:
            st.session_state.option_states = {opt: False for opt in q['options']}
            st.session_state.current_q_id = q['id']

        # 2. 遍历选项并显示 Checkbox
        for opt in q['options']:
            # st.checkbox会返回True/False，我们用它来更新状态
            is_checked = st.checkbox(
                opt, 
                key=f"cb_{q['id']}_{opt}", 
                value=st.session_state.option_states[opt],
                disabled=st.session_state.submitted
            )
            
            # 3. 如果是单选题，并且用户刚刚勾选了这一项
            if q['type'] == '单选' and is_checked and not st.session_state.option_states[opt]:
                # 取消所有其他选项的选中状态
                for other_opt in st.session_state.option_states:
                    st.session_state.option_states[other_opt] = False
                # 只保留当前项为选中
                st.session_state.option_states[opt] = True
                st.rerun() # 立即刷新界面以显示单选效果
            else:
                st.session_state.option_states[opt] = is_checked
        
        # 4. 从状态字典中整理出最终的用户选择
        selected_options = [opt for opt, checked in st.session_state.option_states.items() if checked]
        
        if q['type'] == '单选':
            choice = selected_options[0] if selected_options else None
        else: # 多选
            choice = selected_options

        # --- 修改结束 ---

        if not st.session_state['submitted']:
            if st.button("提交答案", type="primary"):
                if not choice:
                    st.warning("请先选择一个选项！")
                else:
                    st.session_state['submitted'] = True
                    if isinstance(choice, str):
                        user_ans = choice.split('.')[0].strip()
                    else:
                        user_ans = "".join(sorted([c.split('.')[0].strip() for c in choice]))
                    is_correct = (user_ans == q['answer'])
                    st.session_state['last_correct'] = is_correct
                    if is_correct:
                        st.markdown(f'<div class="result-ok">✅ 正确答案：{q["answer"]}</div>', unsafe_allow_html=True)
                    else:
                        st.markdown(f'<div class="result-bad">❌ 正确答案：{q["answer"]}</div>', unsafe_allow_html=True)
                        if q not in st.session_state['mistakes']:
                            st.session_state['mistakes'].append(q)
                    st.rerun()
                    
        else:
            if st.session_state.get('last_correct'):
                st.markdown(f'<div class="result-ok">✅ 正确答案：{q["answer"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="result-bad">❌ 正确答案：{q["answer"]}</div>', unsafe_allow_html=True)
            if st.button("下一题"):
                if st.session_state['current_q'] in st.session_state['quiz_data']:
                    st.session_state['quiz_data'].remove(st.session_state['current_q'])
                st.session_state['current_q'] = None
                st.session_state['submitted'] = False
                st.session_state['last_correct'] = None
                # 清除选项状态以迎接新题目
                if 'option_states' in st.session_state:
                    del st.session_state['option_states']
                st.rerun()

